"""文件处理服务"""

import aiofiles
import os
import time
import gc
import io
import logging
from datetime import datetime
from typing import Optional, List, Dict, Tuple
import PyPDF2
import docx
from fastapi import UploadFile
import aiohttp
import asyncio
from ..config import settings

logger = logging.getLogger(__name__)

# 新增的第三方库
try:
    import pdfplumber
    import fitz  # PyMuPDF
    from docx2python import docx2python
    from PIL import Image

    HAS_ADVANCED_LIBS = True
except ImportError as e:
    HAS_ADVANCED_LIBS = False
    logger.warning("高级文档处理库未安装: %s", e)


class FileService:
    """文件处理服务"""

    ALLOWED_DOCUMENT_TYPES = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    # 图片上传配置
    IMAGE_UPLOAD_URL = "https://mt.agnet.top/image/upload"
    IMAGE_UPLOAD_TIMEOUT = 30  # 超时时间（秒）

    @staticmethod
    def is_supported_document(content_type: str | None) -> bool:
        """判断上传文件类型是否受支持。"""
        return bool(content_type and content_type in FileService.ALLOWED_DOCUMENT_TYPES)

    @staticmethod
    async def upload_image_to_server(image_data: bytes, filename: str) -> Optional[str]:
        """上传图片到外部服务器"""
        try:
            # 准备multipart/form-data格式的数据
            form_data = aiohttp.FormData()
            form_data.add_field(
                "file",
                io.BytesIO(image_data),
                filename=filename,
                content_type="image/jpeg",
            )

            timeout = aiohttp.ClientTimeout(total=FileService.IMAGE_UPLOAD_TIMEOUT)
            async with aiohttp.ClientSession(timeout=timeout) as session:
                async with session.post(
                    FileService.IMAGE_UPLOAD_URL, data=form_data
                ) as response:
                    if response.status == 200:
                        result = await response.json()
                        # 根据实际API返回格式获取图片URL
                        return result.get("file_url")
                    else:
                        logger.warning("图片上传失败，状态码: %s", response.status)
                        return None
        except Exception as e:
            logger.warning("图片上传异常: %s", e)
            return None

    @staticmethod
    def extract_images_from_pdf(file_path: str) -> List[Tuple[bytes, str, int, int]]:
        """从PDF提取图片，返回 (图片数据, 扩展名, 页码, 图片索引) 列表"""
        if not HAS_ADVANCED_LIBS:
            return []

        images = []
        try:
            doc = fitz.open(file_path)

            for page_num in range(doc.page_count):
                page = doc[page_num]
                image_list = page.get_images(full=True)

                for img_index, img in enumerate(image_list):
                    try:
                        # 获取图片数据
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)

                        # 转换为RGB格式（如果是CMYK）
                        if pix.n - pix.alpha < 4:
                            img_data = pix.tobytes("jpeg")
                            ext = "jpg"
                        else:
                            pix1 = fitz.Pixmap(fitz.csRGB, pix)
                            img_data = pix1.tobytes("jpeg")
                            ext = "jpg"
                            pix1 = None

                        pix = None
                        images.append((img_data, ext, page_num + 1, img_index + 1))

                    except Exception as e:
                        logger.warning(
                            "提取PDF第%s页图片%s失败: %s",
                            page_num + 1,
                            img_index + 1,
                            e,
                        )
                        continue

            doc.close()
            return images

        except Exception as e:
            logger.warning("PDF图片提取失败: %s", e)
            return []

    @staticmethod
    def extract_images_from_docx(file_path: str) -> List[Tuple[bytes, str, int]]:
        """从Word文档提取图片，返回 (图片数据, 扩展名, 图片索引) 列表"""
        images = []
        doc = None
        try:
            doc = docx.Document(file_path)

            # 获取文档中的所有关系
            rels = doc.part.rels
            img_index = 0

            for rel in rels.values():
                if "image" in rel.target_ref:
                    try:
                        # 读取图片数据
                        img_data = rel.target_part.blob

                        # 根据content_type确定扩展名
                        content_type = rel.target_part.content_type
                        if "jpeg" in content_type:
                            ext = "jpg"
                        elif "png" in content_type:
                            ext = "png"
                        elif "gif" in content_type:
                            ext = "gif"
                        elif "bmp" in content_type:
                            ext = "bmp"
                        else:
                            ext = "jpg"  # 默认

                        img_index += 1
                        images.append((img_data, ext, img_index))

                    except Exception as e:
                        logger.warning("提取Word文档图片%s失败: %s", img_index + 1, e)
                        continue

            if doc:
                del doc
            gc.collect()
            return images

        except Exception as e:
            if doc:
                del doc
            gc.collect()
            logger.warning("Word文档图片提取失败: %s", e)
            return []

    @staticmethod
    def _safe_file_cleanup(file_path: str, max_retries: int = 3) -> bool:
        """安全删除文件，带重试机制"""
        for attempt in range(max_retries):
            try:
                if os.path.exists(file_path):
                    # 强制垃圾回收，释放可能的文件句柄
                    gc.collect()
                    time.sleep(0.1 * (attempt + 1))  # 递增延迟
                    os.remove(file_path)
                return True
            except OSError as e:
                if attempt == max_retries - 1:
                    logger.warning("无法删除文件 %s: %s", file_path, e)
                    return False
                time.sleep(0.5)  # 等待后重试
        return True

    @staticmethod
    async def save_uploaded_file(file: UploadFile) -> str:
        """保存上传的文件并返回文件路径"""
        # 创建上传目录
        os.makedirs(settings.upload_dir, exist_ok=True)

        # 生成带时间戳的文件名，防止重复
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")[:-3]  # 精确到毫秒
        filename = file.filename or "unknown_file"

        # 分离文件名和扩展名
        name, ext = os.path.splitext(filename)

        # 生成新的文件名：原文件名_时间戳.扩展名
        new_filename = f"{name}_{timestamp}{ext}"
        file_path = os.path.join(settings.upload_dir, new_filename)

        # 异步保存文件
        async with aiofiles.open(file_path, "wb") as f:
            content = await file.read()
            await f.write(content)

        return file_path

    @staticmethod
    async def extract_text_from_pdf(file_path: str) -> str:
        """从PDF文件提取文本，支持表格内容和图片"""
        if HAS_ADVANCED_LIBS:
            return await FileService._extract_pdf_with_pdfplumber(file_path)
        else:
            # 降级到原来的PyPDF2方法
            return FileService._extract_pdf_with_pypdf2(file_path)

    @staticmethod
    async def _extract_pdf_with_pdfplumber(file_path: str) -> str:
        """使用pdfplumber提取PDF文本，包含表格和图片（确保及时释放文件句柄）"""
        try:
            extracted_text = []
            image_references = []  # 存储图片引用映射
            global_img_counter = 1

            # 获取PDF文档的所有图片信息，用于后续匹配
            all_images = FileService.extract_images_from_pdf(file_path)
            page_images_map = {}
            for img_data, ext, page_num, img_index in all_images:
                if page_num not in page_images_map:
                    page_images_map[page_num] = []
                page_images_map[page_num].append((img_data, ext, img_index))

            # 使用上下文管理器，避免在Windows上产生文件锁
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # 添加页码标识
                    extracted_text.append(f"\n--- 第 {page_num} 页 ---\n")

                    # 提取普通文本
                    text = page.extract_text()
                    if text:
                        # 检查文本中是否有图片标记
                        import re

                        img_pattern = r"----.*?(?:image|img|media).*?----"
                        img_matches = list(
                            re.finditer(img_pattern, text, re.IGNORECASE)
                        )

                        if img_matches and page_num in page_images_map:
                            # 按顺序处理页面中的图片
                            page_images = page_images_map[page_num]
                            processed_text = text

                            for i, match in enumerate(img_matches):
                                if i < len(page_images):
                                    # 获取对应的图片数据
                                    img_data, ext, img_index = page_images[i]
                                    filename = (
                                        f"pdf_page{page_num}_img{img_index}.{ext}"
                                    )

                                    # 上传图片
                                    image_url = (
                                        await FileService.upload_image_to_server(
                                            img_data, filename
                                        )
                                    )

                                    if image_url:
                                        # 替换图片标记
                                        old_mark = match.group()
                                        new_mark = f"[图片{global_img_counter}]"
                                        processed_text = processed_text.replace(
                                            old_mark, new_mark, 1
                                        )

                                        # 记录图片引用
                                        image_references.append(
                                            f"[图片{global_img_counter}]: {image_url}"
                                        )
                                        global_img_counter += 1

                            extracted_text.append(processed_text)
                        else:
                            extracted_text.append(text)

                    # 提取表格
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        extracted_text.append(f"\n[表格 {table_num}]")
                        for row in table:
                            if row:  # 跳过空行
                                # 过滤空值并连接单元格
                                row_text = " | ".join(
                                    [str(cell) if cell else "" for cell in row]
                                )
                                extracted_text.append(row_text)
                        extracted_text.append("[表格结束]\n")

            # 在文档末尾添加图片引用映射
            if image_references:
                extracted_text.append(f"\n\n--- 图片引用 ---")
                extracted_text.extend(image_references)

            result = "\n".join(extracted_text).strip()
            gc.collect()
            return result
        except Exception as e:
            gc.collect()
            # 如果pdfplumber失败，尝试PyMuPDF
            try:
                return await FileService._extract_pdf_with_pymupdf(file_path)
            except Exception:
                raise Exception(f"PDF文件读取失败: {str(e)}")

    @staticmethod
    async def _extract_pdf_with_pymupdf(file_path: str) -> str:
        """使用PyMuPDF提取PDF文本和图片"""
        try:
            doc = fitz.open(file_path)
            extracted_text = []

            for page_num in range(doc.page_count):
                page = doc[page_num]
                extracted_text.append(f"\n--- 第 {page_num + 1} 页 ---\n")

                # 提取文本
                text = page.get_text()
                if text:
                    extracted_text.append(text)

                # 尝试提取表格
                try:
                    tables = page.find_tables()
                    for table_num, table in enumerate(tables, 1):
                        extracted_text.append(f"\n[表格 {table_num}]")
                        table_data = table.extract()
                        for row in table_data:
                            if row:
                                row_text = " | ".join(
                                    [str(cell) if cell else "" for cell in row]
                                )
                                extracted_text.append(row_text)
                        extracted_text.append("[表格结束]\n")
                except:
                    # 如果表格提取失败，跳过
                    pass

            doc.close()
            return "\n".join(extracted_text).strip()
        except Exception as e:
            raise Exception(f"PDF文件读取失败: {str(e)}")

    @staticmethod
    def _extract_pdf_with_pypdf2(file_path: str) -> str:
        """使用PyPDF2提取PDF文本（原方法）"""
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            raise Exception(f"PDF文件读取失败: {str(e)}")

    @staticmethod
    async def extract_text_from_docx(file_path: str) -> str:
        """从Word文档提取文本，支持表格内容和图片"""
        if HAS_ADVANCED_LIBS:
            return await FileService._extract_docx_with_docx2python(file_path)
        else:
            # 降级到原来的python-docx方法，但增强表格处理
            return await FileService._extract_docx_with_python_docx(file_path)

    @staticmethod
    async def _extract_docx_with_docx2python(file_path: str) -> str:
        """使用docx2python提取Word文档内容和图片（确保及时释放文件句柄）"""
        try:
            extracted_text = []
            image_references = []  # 存储图片引用映射
            global_img_counter = 1

            # 获取Word文档的所有图片信息
            all_images = FileService.extract_images_from_docx(file_path)

            # 使用上下文管理器确保文件及时关闭，避免Windows上的锁定
            with docx2python(file_path) as content:
                # 处理文档内容
                if hasattr(content, "document"):
                    for section in content.document:
                        for element in section:
                            if isinstance(element, list):
                                # 这可能是表格
                                extracted_text.append("\n[表格内容]")
                                for row in element:
                                    if isinstance(row, list):
                                        row_text = " | ".join(
                                            [str(cell).strip() for cell in row if cell]
                                        )
                                        if row_text:
                                            extracted_text.append(row_text)
                                    else:
                                        extracted_text.append(str(row))
                                extracted_text.append("[表格结束]\n")
                            else:
                                # 普通文本，检查是否包含图片标记
                                text = str(element).strip()
                                if text:
                                    # 检查文本中是否有图片标记
                                    import re

                                    img_pattern = r"----.*?(?:image|img|media).*?----"
                                    img_matches = list(
                                        re.finditer(img_pattern, text, re.IGNORECASE)
                                    )

                                    if img_matches and all_images:
                                        processed_text = text

                                        for match in img_matches:
                                            if global_img_counter <= len(all_images):
                                                # 获取对应的图片数据
                                                img_data, ext, img_index = all_images[
                                                    global_img_counter - 1
                                                ]
                                                filename = f"docx_img{global_img_counter}.{ext}"

                                                # 上传图片
                                                image_url = await FileService.upload_image_to_server(
                                                    img_data, filename
                                                )

                                                if image_url:
                                                    # 替换图片标记
                                                    old_mark = match.group()
                                                    new_mark = (
                                                        f"[图片{global_img_counter}]"
                                                    )
                                                    processed_text = (
                                                        processed_text.replace(
                                                            old_mark, new_mark, 1
                                                        )
                                                    )

                                                    # 记录图片引用
                                                    image_references.append(
                                                        f"[图片{global_img_counter}]: {image_url}"
                                                    )
                                                    global_img_counter += 1

                                        extracted_text.append(processed_text)
                                    else:
                                        extracted_text.append(text)

            # 在文档末尾添加图片引用映射
            if image_references:
                extracted_text.append(f"\n\n--- 图片引用 ---")
                extracted_text.extend(image_references)

            result = "\n".join(extracted_text).strip()
            gc.collect()
            return result
        except Exception as e:
            gc.collect()
            # 如果docx2python失败，回退到增强的python-docx
            try:
                return await FileService._extract_docx_with_python_docx(file_path)
            except Exception:
                raise Exception(f"Word文档读取失败: {str(e)}")

    @staticmethod
    async def _extract_docx_with_python_docx(file_path: str) -> str:
        """使用python-docx提取Word文档内容和图片（增强版）"""
        doc = None
        try:
            doc = docx.Document(file_path)
            extracted_text = []
            image_references = []  # 存储图片引用映射
            global_img_counter = 1

            # 获取Word文档的所有图片信息
            all_images = FileService.extract_images_from_docx(file_path)

            # 提取段落文本，同时处理图片
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # 检查文本中是否有图片标记
                    import re

                    img_pattern = r"----.*?(?:image|img|media).*?----"
                    img_matches = list(re.finditer(img_pattern, text, re.IGNORECASE))

                    if img_matches and all_images:
                        processed_text = text

                        for match in img_matches:
                            if global_img_counter <= len(all_images):
                                # 获取对应的图片数据
                                img_data, ext, img_index = all_images[
                                    global_img_counter - 1
                                ]
                                filename = f"docx_img{global_img_counter}.{ext}"

                                # 上传图片
                                image_url = await FileService.upload_image_to_server(
                                    img_data, filename
                                )

                                if image_url:
                                    # 替换图片标记
                                    old_mark = match.group()
                                    new_mark = f"[图片{global_img_counter}]"
                                    processed_text = processed_text.replace(
                                        old_mark, new_mark, 1
                                    )

                                    # 记录图片引用
                                    image_references.append(
                                        f"[图片{global_img_counter}]: {image_url}"
                                    )
                                    global_img_counter += 1

                        extracted_text.append(processed_text)
                    else:
                        extracted_text.append(text)

            # 提取表格内容
            for table_num, table in enumerate(doc.tables, 1):
                extracted_text.append(f"\n[表格 {table_num}]")
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text if cell_text else "")
                    row_text = " | ".join(row_data)
                    if row_text.strip():
                        extracted_text.append(row_text)
                extracted_text.append("[表格结束]\n")

            # 在文档末尾添加图片引用映射
            if image_references:
                extracted_text.append(f"\n\n--- 图片引用 ---")
                extracted_text.extend(image_references)

            result = "\n".join(extracted_text).strip()

            # 确保释放资源
            if doc:
                del doc
            gc.collect()

            return result
        except Exception as e:
            # 确保释放资源
            if doc:
                del doc
            gc.collect()
            raise Exception(f"Word文档读取失败: {str(e)}")

    @staticmethod
    async def process_uploaded_file(file: UploadFile) -> str:
        """处理上传的文件并提取文本内容"""
        # 检查文件大小
        content = await file.read()
        if len(content) > settings.max_file_size:
            raise Exception(
                f"文件大小超过限制 ({settings.max_file_size / 1024 / 1024}MB)"
            )

        # 重置文件指针
        await file.seek(0)

        # 保存文件
        file_path = await FileService.save_uploaded_file(file)

        try:
            # 根据文件类型提取文本和图片
            if file.content_type == "application/pdf":
                text = await FileService.extract_text_from_pdf(file_path)
            elif (
                file.content_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                text = await FileService.extract_text_from_docx(file_path)
            else:
                raise Exception("不支持的文件类型，请上传PDF或Word文档")

            # 成功提取后，使用安全的文件清理方法
            FileService._safe_file_cleanup(file_path)

            return text

        except Exception as e:
            # 异常情况下也使用安全的文件清理方法
            FileService._safe_file_cleanup(file_path)
            raise e
