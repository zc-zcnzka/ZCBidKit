"""Word 导出服务。"""

import io
import re
from urllib.parse import quote

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from ..models.schemas import WordExportOutlineItem, WordExportRequest


def _set_run_font_simsun(run: docx.text.run.Run) -> None:
    run.font.name = "宋体"
    rpr = run._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        rpr.rFonts.set(qn("w:eastAsia"), "宋体")


def _set_paragraph_font_simsun(paragraph: docx.text.paragraph.Paragraph) -> None:
    for run in paragraph.runs:
        _set_run_font_simsun(run)


class WordExportService:
    """负责将目录数据导出为 Word 文档。"""

    @staticmethod
    def export_outline(request: WordExportRequest) -> tuple[io.BytesIO, dict[str, str]]:
        doc = docx.Document()
        WordExportService._init_document_styles(doc)
        WordExportService._add_document_intro(doc, request.project_name)
        WordExportService._add_outline_items(doc, request.outline)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{request.project_name or '标书文档'}.docx"
        headers = {
            "Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}",
        }
        return buffer, headers

    @staticmethod
    def _init_document_styles(doc: docx.Document) -> None:
        try:
            styles = doc.styles
            base_styles = ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title"]
            for style_name in base_styles:
                if style_name not in styles:
                    continue
                style = styles[style_name]
                font = style.font
                font.name = "宋体"
                if style._element.rPr is None:
                    style._element._add_rPr()
                style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                if style_name == "Normal":
                    font.bold = False
        except Exception:
            pass

    @staticmethod
    def _add_document_intro(
        doc: docx.Document, project_name: str | None
    ) -> None:
        declaration = doc.add_paragraph()
        declaration_run = declaration.add_run("内容由AI生成")
        declaration_run.italic = True
        declaration_run.font.size = Pt(9)
        _set_run_font_simsun(declaration_run)
        declaration.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(project_name or "投标技术文件")
        title_run.bold = True
        title_run.font.size = Pt(16)
        _set_run_font_simsun(title_run)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    @staticmethod
    def _add_markdown_runs(para: docx.text.paragraph.Paragraph, text: str) -> None:
        pattern = r"(\*\*.*?\*\*|\*.*?\*|`.*?`)"
        parts = re.split(pattern, text)
        for part in parts:
            if not part:
                continue
            run = para.add_run()
            if part.startswith("**") and part.endswith("**") and len(part) > 4:
                run.text = part[2:-2]
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                run.text = part[1:-1]
                run.italic = True
            elif part.startswith("`") and part.endswith("`") and len(part) > 2:
                run.text = part[1:-1]
            else:
                run.text = part
            _set_run_font_simsun(run)

    @staticmethod
    def _add_markdown_paragraph(doc: docx.Document, text: str) -> None:
        para = doc.add_paragraph()
        WordExportService._add_markdown_runs(para, text)
        para.paragraph_format.space_after = Pt(6)

    @staticmethod
    def _parse_markdown_blocks(content: str) -> list[tuple]:
        blocks: list[tuple] = []
        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i].rstrip("\r").strip()
            if not line:
                i += 1
                continue

            if (
                line.startswith("- ")
                or line.startswith("* ")
                or re.match(r"^\d+\.\s", line)
            ):
                items: list[tuple] = []
                while i < len(lines):
                    raw = lines[i].rstrip("\r")
                    stripped = raw.strip()
                    if stripped.startswith("- ") or stripped.startswith("* "):
                        text = re.sub(r"^[-*]\s+", "", stripped).strip()
                        if text:
                            items.append(("unordered", None, text))
                        i += 1
                        continue
                    match_number = re.match(r"^(\d+)\.\s+(.*)$", stripped)
                    if match_number:
                        num_str, text = match_number.groups()
                        if text.strip():
                            items.append(("ordered", num_str, text.strip()))
                        i += 1
                        continue
                    break
                if items:
                    blocks.append(("list", items))
                continue

            if "|" in line:
                rows: list[str] = []
                while i < len(lines):
                    stripped = lines[i].rstrip("\r").strip()
                    if "|" not in stripped:
                        break
                    if not re.match(r"^\|?[-\s\|]+\|?$", stripped):
                        cells = [cell.strip() for cell in stripped.split("|")]
                        row_text = " | ".join([cell for cell in cells if cell])
                        if row_text:
                            rows.append(row_text)
                    i += 1
                if rows:
                    blocks.append(("table", rows))
                continue

            if line.startswith("#"):
                match_heading = re.match(r"^(#+)\s*(.*)$", line)
                if match_heading:
                    level_marks, title_text = match_heading.groups()
                    blocks.append(
                        ("heading", min(len(level_marks), 3), title_text.strip())
                    )
                i += 1
                continue

            para_lines: list[str] = []
            while i < len(lines):
                stripped = lines[i].rstrip("\r").strip()
                if (
                    stripped
                    and not stripped.startswith("-")
                    and not stripped.startswith("*")
                    and "|" not in stripped
                    and not stripped.startswith("#")
                ):
                    para_lines.append(stripped)
                    i += 1
                else:
                    break
            if para_lines:
                blocks.append(("paragraph", " ".join(para_lines)))
            else:
                i += 1

        return blocks

    @staticmethod
    def _render_markdown_blocks(doc: docx.Document, blocks: list[tuple]) -> None:
        for block in blocks:
            kind = block[0]
            if kind == "list":
                for item_kind, num_str, text in block[1]:
                    paragraph = doc.add_paragraph()
                    prefix = "• " if item_kind == "unordered" else f"{num_str}. "
                    run = paragraph.add_run(prefix)
                    _set_run_font_simsun(run)
                    WordExportService._add_markdown_runs(paragraph, text)
            elif kind == "table":
                for row in block[1]:
                    WordExportService._add_markdown_paragraph(doc, row)
            elif kind == "heading":
                _, level, text = block
                heading = doc.add_heading(text, level=level)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _set_paragraph_font_simsun(heading)
            elif kind == "paragraph":
                WordExportService._add_markdown_paragraph(doc, block[1])

    @staticmethod
    def _add_markdown_content(doc: docx.Document, content: str) -> None:
        blocks = WordExportService._parse_markdown_blocks(content)
        WordExportService._render_markdown_blocks(doc, blocks)

    @staticmethod
    def _add_outline_items(
        doc: docx.Document, items: list[WordExportOutlineItem], level: int = 1
    ) -> None:
        for item in items:
            if level <= 3:
                heading = doc.add_heading(f"{item.id} {item.title}", level=level)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _set_paragraph_font_simsun(heading)
            else:
                para = doc.add_paragraph()
                run = para.add_run(f"{item.id} {item.title}")
                run.bold = True
                _set_run_font_simsun(run)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(3)

            if not item.children:
                content = item.content or ""
                if content.strip():
                    WordExportService._add_markdown_content(doc, content)
                continue

            WordExportService._add_outline_items(doc, item.children, level + 1)
